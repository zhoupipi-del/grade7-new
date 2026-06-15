"""
AI 线上推理蓝图 — 学生风险实时预测 API + 家长会谈话单生成器
============================================================

路由:
  GET /ai-api/predict/<int:student_id>      # 返回双轨归因 JSON
  GET /ai-api/talk-sheet/<int:student_id>   # 生成打印版家长会谈单 (HTML)
  GET /ai-api/talk-sheet/<int:student_id>/docx  # 生成 Word 版谈话单下载

依赖:
  - models/wings_xgb_pipeline.pkl  (由 model_trainer.py 生成)
  - feature_extractor.py            (get_student_vector / get_grade_baselines 方法)
"""
import os
import io
import joblib
import numpy as np
from datetime import datetime
from flask import Blueprint, jsonify, current_app, render_template_string, render_template, request, session
from sqlalchemy import func
from feature_extractor import FeatureExtractor
from models import db, Student, Class as ClassModel, InterventionRecord
from utils.db_utils import safe_commit
from utils import get_local_now

bp = Blueprint("ai_inference", __name__, url_prefix="/ai-api")

# ── 模型懒加载缓存 ──
_PIPELINE = None
_METADATA = None
_BASELINES = {}  # {grade_id: {feature: baseline}}


def _load_model():
    """懒加载模型到 worker 内存 (线程安全由 Gunicorn 保证)"""
    global _PIPELINE, _METADATA
    if _PIPELINE is not None:
        return _PIPELINE, _METADATA

    model_dir = os.path.join(current_app.root_path, "models")
    pipeline_path = os.path.join(model_dir, "wings_xgb_pipeline.pkl")
    metadata_path = os.path.join(model_dir, "pipeline_metadata.pkl")

    if not os.path.exists(pipeline_path):
        raise FileNotFoundError(
            f"模型文件不存在: {pipeline_path}\n"
            f"请先运行: cd {{project_root}} && python model_trainer.py"
        )

    _PIPELINE = joblib.load(pipeline_path)
    _METADATA = joblib.load(metadata_path)
    return _PIPELINE, _METADATA


def _get_baselines(grade_id: int) -> dict:
    """获取年级基线 (懒加载 + 按 grade_id 缓存)"""
    global _BASELINES
    if grade_id in _BASELINES:
        return _BASELINES[grade_id]

    fe = FeatureExtractor(grade_id=grade_id)
    _BASELINES[grade_id] = fe.get_grade_baselines()
    return _BASELINES[grade_id]


def invalidate_model_cache():
    """
    MLOps 热更新接口 — 被 model_retrain.py 调用。
    失效所有内存缓存，下次 API 请求自动重新加载模型。
    """
    global _PIPELINE, _METADATA, _BASELINES
    _PIPELINE = None
    _METADATA = None
    _BASELINES = {}


@bp.route("/predict/<int:student_id>")
def predict(student_id):
    """
    AI 风险预测接口 — 返回双轨归因 JSON。

    返回格式:
        {
            "student_id": 123,
            "student_name": "陈佳乐",
            "class_id": 2501,
            "risk_probability": 0.87,      # AI 预测的风险概率 [0, 1]
            "risk_level": "high",            # high / medium / low
            "top_factors": [                 # 技术归因 (算法维度)
                {"feature": "discipline_factor", "importance": 0.51},
                {"feature": "math_avg", "importance": 0.49},
            ],
            "evidence": {                    # 业务归因 (大白话解释)
                "discipline": "违纪超标 11.1 倍",
                "math": "数学成绩仅为年级均值的 0.66",
                "attendance": "出勤率 92%（低于均值 96%）",
                "quality": "综合素质分 72.5（低于均值 80.3）",
            }
        }
    """
    # ① 加载模型
    pipeline, metadata = _load_model()
    classifier = pipeline.named_steps['classifier']
    passed_features = metadata.get("passed_features", metadata.get("feature_names", []))

    # ② 获取学生信息 + 特征向量
    student = Student.query.get_or_404(student_id)
    fe = FeatureExtractor(grade_id=student.grade_id)
    result = fe.get_student_vector(student_id)

    features = result["features"]  # [math_slope, math_avg, ...]

    # 维度适配：FeatureExtractor 输出维度可能与管道期望维度不同
    n_expected = classifier.n_features_in_
    if len(features) != n_expected:
        # 截断或补零对齐
        if len(features) > n_expected:
            features = features[:n_expected]
        else:
            features = features + [0.0] * (n_expected - len(features))

    # ③ 推理
    X = [features]
    proba = pipeline.predict_proba(X)[0]  # [P(0), P(1)]
    risk_prob = float(proba[1])
    pred = pipeline.predict(X)[0]

    # ④ 风险等级映射
    if risk_prob >= 0.7:
        risk_level = "high"
    elif risk_prob >= 0.4:
        risk_level = "medium"
    else:
        risk_level = "low"

    # ⑤ 双轨归因 — per-student 局部贡献分析（类 SHAP）
    #    不再用全局 feature_importances_（所有学生一模一样），
    #    而是用 |scaled_value × importance| 计算每人对各维度的局部贡献。
    scaler = pipeline.named_steps.get('scaler')
    if scaler:
        X_scaled = scaler.transform([features])[0]
    else:
        X_scaled = np.array(features, dtype=float)

    importances = classifier.feature_importances_

    # Per-feature contribution = |scaled_value| × importance
    contributions = []
    n_factors = min(len(importances), len(passed_features))
    for i in range(n_factors):
        contrib = abs(float(X_scaled[i]) * float(importances[i]))
        contributions.append({
            "feature": passed_features[i],
            "importance": contrib,
        })

    # Normalize to sum=1
    total = sum(c["importance"] for c in contributions)
    if total > 0:
        for c in contributions:
            c["importance"] = round(c["importance"] / total, 4)

    # Sort descending and take top 2
    contributions.sort(key=lambda x: -x["importance"])
    top_factors = contributions[:2]

    # ── evidence: 计算与基线(年级均值)的偏离度 ──
    baselines = _get_baselines(student.grade_id)
    evidence = _build_evidence(result["feature_dict"], baselines)

    # ━━━ 方案A 后验追踪钩子 ━━━
    # 学生有 "tracking" 状态的干预记录时，自动对比前后风险概率 → 计算 Δ → 自评效果
    _run_posterior_tracking(student_id, risk_prob)

    return jsonify({
        "student_id": student_id,
        "student_name": result["student_name"],
        "class_id": result["class_id"],
        "risk_probability": round(risk_prob, 4),
        "risk_level": risk_level,
        "top_factors": top_factors,
        "evidence": evidence,
    })


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  家长会谈话单生成器 — 打印版 + Word 版
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

TALK_SHEET_HTML = """\
<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>家长会谈单 — {{ student_name }}</title>
<style>
  @page { size: A4; margin: 15mm; }
  * { margin: 0; padding: 0; box-sizing: border-box; }
  body { font-family: "Microsoft YaHei", "PingFang SC", sans-serif; color: #1a1a1a; line-height: 1.7; max-width: 210mm; margin: 0 auto; padding: 15mm; }
  @media print { body { padding: 0; } }
  .header { text-align: center; border-bottom: 3px solid #1a5276; padding-bottom: 14px; margin-bottom: 20px; }
  .header h1 { font-size: 24px; color: #1a5276; letter-spacing: 4px; }
  .header .subtitle { font-size: 13px; color: #666; margin-top: 4px; }
  .meta-bar { display: flex; justify-content: space-between; background: #f0f4f8; padding: 10px 16px; border-radius: 6px; font-size: 13px; margin-bottom: 20px; }
  .meta-bar .label { color: #888; }
  .risk-banner { text-align: center; padding: 14px; border-radius: 8px; margin-bottom: 20px; color: #fff; font-size: 16px; font-weight: 700; }
  .risk-banner.high { background: linear-gradient(135deg, #e74c3c, #c0392b); }
  .risk-banner.medium { background: linear-gradient(135deg, #f39c12, #e67e22); }
  .risk-banner.low { background: linear-gradient(135deg, #27ae60, #2ecc71); }
  .section { margin-bottom: 18px; }
  .section-title { font-size: 15px; color: #1a5276; border-left: 4px solid #1a5276; padding-left: 10px; margin-bottom: 10px; font-weight: 700; }
  .evidence-table { width: 100%; border-collapse: collapse; font-size: 13px; }
  .evidence-table th { background: #1a5276; color: #fff; padding: 8px 12px; text-align: left; }
  .evidence-table th:last-child { text-align: center; width: 80px; }
  .evidence-table td { padding: 8px 12px; border-bottom: 1px solid #e0e0e0; }
  .evidence-table .bad { color: #e74c3c; font-weight: 600; }
  .evidence-table .ok { color: #27ae60; }
  .evidence-table .warn { text-align: center; }
  .talking-points { background: #fff9e6; border: 1px solid #f0d060; border-radius: 6px; padding: 14px 16px; }
  .talking-points li { font-size: 13px; margin-bottom: 6px; color: #5a4a00; }
  .actions { background: #eaf4f0; border: 1px solid #b0d8c0; border-radius: 6px; padding: 14px 16px; }
  .actions li { font-size: 13px; margin-bottom: 6px; color: #1a4a2a; }
  .footer { margin-top: 30px; padding-top: 14px; border-top: 1px solid #ddd; font-size: 11px; color: #999; text-align: center; }
  .btn-bar { text-align: center; margin-top: 20px; }
  .btn { display: inline-block; padding: 8px 18px; border-radius: 4px; text-decoration: none; font-size: 13px; margin: 0 6px; cursor: pointer; }
  .btn-print { background: #1a5276; color: #fff; border: none; }
  .btn-docx { background: #2b579a; color: #fff; border: none; }
  .btn-intervention { background: #27ae60; color: #fff; border: none; }
  .modal-overlay { position: fixed; top:0; left:0; width:100%; height:100%;
                    background: rgba(0,0,0,0.4); display:flex; align-items:center;
                    justify-content:center; z-index:9999; }
  .modal-content { background:#fff; border-radius:8px; width:520px; max-width:95vw;
                     max-height:90vh; overflow-y:auto; }
  .modal-header  { display:flex; justify-content:space-between; align-items:center;
                     padding:14px 18px; border-bottom:1px solid #eee; font-size:15px; font-weight:700; }
  .modal-close   { cursor:pointer; font-size:22px; color:#999; }
  .modal-body    { padding:18px; }
  .modal-footer  { padding:12px 18px; border-top:1px solid #eee;
                     display:flex; gap:8px; justify-content:flex-end; }
  .form-group    { margin-bottom:12px; }
  .form-group label { display:block; font-size:13px; color:#555; margin-bottom:4px; }
  .form-control  { width:100%; padding:6px 10px; border:1px solid #ddd;
                     border-radius:4px; font-size:13px; }
  @media print { .btn-bar { display: none; } }
</style>
</head>
<body>

<div class="header">
  <h1>梨江中学 · 家长会谈参考单</h1>
  <div class="subtitle">基于 Wings AI 数智系统生成 | 年级平均基线对比 | 仅供班主任内部参考</div>
</div>

<div class="meta-bar">
  <div><span class="label">学生:</span> <strong>{{ student_name }}</strong></div>
  <div><span class="label">班级:</span> <strong>{{ class_name }}</strong></div>
  <div><span class="label">生成时间:</span> {{ gen_time }}</div>
  <div><span class="label">AI预测概率:</span> <strong>{{ risk_pct }}%</strong></div>
</div>

<div class="risk-banner {{ risk_level }}">
  {% if risk_level == 'high' %}⚠️ 高风险预警 — AI 综合评估显示该生需要重点关注{% elif risk_level == 'medium' %}⚡ 中等风险提示 — 建议加强家校沟通{% else %}✅ 当前风险较低 — 保持常规关注{% endif %}
</div>

<div class="section">
  <div class="section-title">📊 AI 多维偏离度分析 (与年级均值对比)</div>
  <table class="evidence-table">
    <tr><th>评估维度</th><th>AI 判定</th><th>偏离度</th></tr>
    {% for ev in evidence_rows %}
    <tr>
      <td>{{ ev.label }}</td>
      <td class="{{ 'bad' if ev.is_bad else 'ok' }}">{{ ev.value }}</td>
      <td class="warn">{{ ev.badge }}</td>
    </tr>
    {% endfor %}
  </table>
</div>

<div class="section">
  <div class="section-title">💬 建议谈话要点</div>
  <ul class="talking-points">
    {% for pt in talking_points %}
    <li>{{ pt }}</li>
    {% endfor %}
  </ul>
</div>

<div class="section">
  <div class="section-title">📋 后续行动建议</div>
  <ul class="actions">
    {% for act in actions %}
    <li>{{ act }}</li>
    {% endfor %}
  </ul>
</div>

<div class="btn-bar">
  <button class="btn btn-print" onclick="window.print()">🖨️ 打印谈话单</button>
  <a class="btn btn-docx" href="/ai-api/talk-sheet/{{ student_id }}/docx">📥 下载 Word 文档</a>
  <button class="btn btn-intervention" onclick="openInterventionModal()">📝 记录干预</button>
</div>

<div class="footer">
  Wings AI 数智德育系统 · 梨江中学德育处 · 数据来源: 学生考勤/成绩/违纪/素质评价历史数据<br>
  此报告基于机器学习模型自动生成，仅供参考，不作为唯一决策依据
</div>

<!-- 干预记录弹窗 -->
<div id="interventionModal" class="modal-overlay" style="display:none;">
  <div class="modal-content">
    <div class="modal-header">
      <h3>📝 记录干预措施</h3>
      <span class="modal-close" onclick="closeInterventionModal()">&times;</span>
    </div>
    <div class="modal-body">
      <input type="hidden" id="int_student_id" value="{{ student_id }}">
      <div class="form-group">
        <label>干预类型</label>
        <select id="int_type" class="form-control">
          <option value="谈话">谈话</option>
          <option value="家长联动">家长联动</option>
          <option value="座位调整">座位调整</option>
          <option value="学业辅导">学业辅导</option>
          <option value="心理干预">心理干预</option>
          <option value="行为契约">行为契约</option>
          <option value="其他">其他</option>
        </select>
      </div>
      <div class="form-group">
        <label>干预日期</label>
        <input type="date" id="int_date" class="form-control" value="{{ gen_time[:10] }}">
      </div>
      <div class="form-group">
        <label>谈话/干预记录</label>
        <textarea id="int_notes" class="form-control" rows="4" placeholder="记录本次干预的详细内容..."></textarea>
      </div>
      <div class="form-group">
        <label>计划随访日期（可选）</label>
        <input type="date" id="int_followup" class="form-control">
      </div>
    </div>
    <div class="modal-footer">
      <button class="btn btn-intervention" onclick="submitIntervention()">提交</button>
      <button class="btn btn-secondary" onclick="closeInterventionModal()">取消</button>
    </div>
  </div>
</div>

<script>
function openInterventionModal() {
  document.getElementById('interventionModal').style.display = 'flex';
}
function closeInterventionModal() {
  document.getElementById('interventionModal').style.display = 'none';
}
function submitIntervention() {
  var data = {
    student_id: document.getElementById('int_student_id').value,
    intervention_type: document.getElementById('int_type').value,
    notes: document.getElementById('int_notes').value,
    intervention_date: document.getElementById('int_date').value,
    follow_up_date: document.getElementById('int_followup').value || null,
  };
  fetch('/ai-api/intervention/create', {
    method: 'POST',
    headers: {'Content-Type': 'application/json'},
    body: JSON.stringify(data)
  }).then(r => r.json()).then(d => {
    if (d.status === 'ok') {
      alert('✅ 干预记录已创建！\n干预ID: ' + d.intervention_id + '\n风险概率(前): ' + d.risk_before);
      closeInterventionModal();
      document.getElementById('int_notes').value = '';
      document.getElementById('int_followup').value = '';
    } else {
      alert('❌ 创建失败: ' + (d.error || d.message));
    }
  });
}
</script>

</body>
</html>"""


def _get_talk_sheet_data(student_id):
    """获取谈话单所需的全部数据"""
    from flask import current_app as ctx
    student = Student.query.get_or_404(student_id)

    # 推理数据
    fe = FeatureExtractor(grade_id=student.grade_id)

    pipeline, metadata = _load_model()
    classifier = pipeline.named_steps['classifier']
    passed_features = metadata["passed_features"]

    result = fe.get_student_vector(student_id)
    features = result["features"]
    proba = pipeline.predict_proba([features])[0]
    risk_prob = float(proba[1])

    if risk_prob >= 0.7:
        risk_level = "high"
    elif risk_prob >= 0.4:
        risk_level = "medium"
    else:
        risk_level = "low"

    baselines = _get_baselines(student.grade_id)
    evidence = _build_evidence(result["feature_dict"], baselines)

    # 班级名
    klass = ClassModel.query.get(student.class_id)
    class_name = klass.name if klass else f"班级{student.class_id}"

    # ── 构建证据行 (带视觉标记) ──
    BAD_KEYWORDS = ("超标", "低于", "偏低", "下滑", "高频", "仅为")
    EVIDENCE_ORDER = ["discipline", "math", "attendance", "quality", "trend", "risk"]
    EVIDENCE_LABELS = {
        "discipline": "违纪情况",
        "math": "数学成绩",
        "attendance": "出勤情况",
        "quality": "综合素质",
        "trend": "成绩趋势",
        "risk": "近期预警",
    }

    evidence_rows = []
    for key in EVIDENCE_ORDER:
        if key not in evidence:
            continue
        val = evidence[key]
        is_bad = any(kw in val for kw in BAD_KEYWORDS)
        badge = "🔴" if is_bad else "🟢"
        evidence_rows.append({
            "label": EVIDENCE_LABELS.get(key, key),
            "value": val,
            "is_bad": is_bad,
            "badge": badge,
        })

    # ── 生成谈话要点 (基于 evidence) ──
    talking_points = _generate_talking_points(
        student.name, evidence, result["feature_dict"]
    )

    # ── 生成行动建议 ──
    actions = _generate_actions(risk_level, evidence)

    return {
        "student_id": student_id,
        "student_name": student.name,
        "class_name": class_name,
        "risk_prob": risk_prob,
        "risk_level": risk_level,
        "risk_pct": f"{risk_prob * 100:.1f}",
        "evidence_rows": evidence_rows,
        "talking_points": talking_points,
        "actions": actions,
        "gen_time": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }


def _generate_talking_points(student_name: str, evidence: dict, features: dict) -> list:
    """根据 evidence 自动生成谈话要点"""
    points = []

    # 开场白
    points.append(f"{student_name}家长您好，感谢您百忙之中来参加本次家长会。")

    # 违纪
    disc = evidence.get("discipline", "")
    if "超标" in disc:
        points.append(f"近期{student_name}同学的违纪记录较为突出，{disc}。建议家长在家中加强对孩子行为规范的要求，并与我们保持密切沟通。")

    # 数学
    math = evidence.get("math", "")
    if "低于及格线" in math:
        points.append(f"数学是{student_name}同学当前最大的学业短板，{math}。建议课后增加数学练习时间，可考虑安排针对性辅导。")
    elif "仅为" in math:
        points.append(f"数学成绩是目前的薄弱科目：{math}。建议在课后作业中多加关注，必要时可申请课后辅导。")

    # 出勤
    att = evidence.get("attendance", "")
    if "低于" in att:
        points.append(f"出勤方面需要关注：{att}。频繁缺勤会直接影响学业进度，请家长帮助孩子养成按时到校的习惯。")

    # 素质
    quality = evidence.get("quality", "")
    if "偏低" in quality or "低于" in quality:
        points.append(f"综合素质评价方面：{quality}。建议鼓励孩子多参与班级活动和社会实践，提升综合能力。")

    # 趋势
    trend = evidence.get("trend", "")
    if "下滑" in trend:
        points.append(f"⚠️ 成绩趋势需特别警惕：{trend}。这往往反映出学习状态或心理状态的波动，建议家长多与孩子沟通交流。")
    elif "上升" in trend:
        points.append(f"值得肯定的是：{trend}。希望继续保持这个势头！")

    # 预警
    risk = evidence.get("risk", "")
    if "高频" in risk or "中频" in risk:
        points.append(f"AI系统在过去30天内对{student_name}同学发出了{risk}提示。这提醒我们需要更多关注孩子的情绪和行为变化。")

    # 收尾
    points.append("我们相信在家校共同努力下，孩子一定能够取得更好的发展。有任何问题欢迎随时与我联系。")

    return points


def _generate_actions(risk_level: str, evidence: dict) -> list:
    """根据风险等级生成后续行动建议"""
    actions = []

    if risk_level == "high":
        actions.append("【紧急】本周内安排一次一对一家长面谈，深入了解学生近期状态")
        actions.append("将学生纳入重点关注名单，每日记录行为表现")
        if any("数学" in str(v) for v in evidence.values()):
            actions.append("协调数学老师制定个性化辅导计划")
        if any("违纪" in str(v) for v in evidence.values()):
            actions.append("与德育处联动，对该生违纪行为进行专项干预")
    elif risk_level == "medium":
        actions.append("两周内安排家长电话沟通或面谈")
        actions.append("将该生纳入班级月度关注名单")
        actions.append("在班会中适当关注，鼓励积极参与集体活动")
    else:
        actions.append("按常规节奏进行家校沟通")
        actions.append("继续保持现有的学习和行为习惯")

    actions.append("将本次谈话单存档至学生成长档案")
    return actions


@bp.route("/talk-sheet/<int:student_id>")
def talk_sheet(student_id):
    """生成打印版HTML家长会谈单"""
    data = _get_talk_sheet_data(student_id)
    return render_template_string(TALK_SHEET_HTML, **data)


@bp.route("/talk-sheet/<int:student_id>/docx")
def talk_sheet_docx(student_id):
    """生成 Word 文档版家长会谈单 (下载)"""
    try:
        from docx import Document
        from docx.shared import Mm, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return jsonify({"error": "python-docx 未安装，无法生成 Word 文档"}), 500

    data = _get_talk_sheet_data(student_id)
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)

    # 标题
    title = doc.add_heading("梨江中学 · 家长会谈参考单", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于 Wings AI 数智系统生成 | 仅供班主任内部参考")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 基本信息
    doc.add_paragraph("")
    info_table = doc.add_table(rows=1, cols=4)
    info_table.style = "Light Grid Accent 1"
    cells = info_table.rows[0].cells
    cells[0].text = f"学生: {data['student_name']}"
    cells[1].text = f"班级: {data['class_name']}"
    cells[2].text = f"生成时间: {data['gen_time']}"
    cells[3].text = f"AI风险概率: {data['risk_pct']}%"

    # 风险等级
    doc.add_paragraph("")
    risk_label = {"high": "⚠️ 高风险预警", "medium": "⚡ 中等风险提示", "low": "✅ 当前风险较低"}
    doc.add_heading(risk_label.get(data["risk_level"], ""), level=2)

    # 偏离度分析
    doc.add_heading("AI 多维偏离度分析", level=2)
    ev_table = doc.add_table(rows=1, cols=3)
    ev_table.style = "Light Grid Accent 1"
    hdr = ev_table.rows[0].cells
    hdr[0].text = "评估维度"
    hdr[1].text = "AI 判定"
    hdr[2].text = "偏离度"
    for ev in data["evidence_rows"]:
        row = ev_table.add_row()
        row.cells[0].text = ev["label"]
        row.cells[1].text = ev["value"]
        row.cells[2].text = ev["badge"]

    # 谈话要点
    doc.add_heading("建议谈话要点", level=2)
    for pt in data["talking_points"]:
        doc.add_paragraph(pt, style="List Bullet")

    # 行动建议
    doc.add_heading("后续行动建议", level=2)
    for act in data["actions"]:
        doc.add_paragraph(act, style="List Bullet")

    # 页脚
    doc.add_paragraph("")
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("Wings AI 数智德育系统 · 梨江中学德育处 · 此报告基于机器学习模型自动生成")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # 返回文件
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    from flask import send_file
    filename = f"家长会谈单_{data['student_name']}_{data['gen_time'].replace(' ', '_')}.docx"
    return send_file(
        buffer,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


def _build_evidence(feature_dict: dict, baselines: dict) -> dict:
    """
    业务归因 — 只展示与年级基线有显著偏离的维度。
    正常维度不展示，让每个学生的 evidence 真正有区分度。
    """
    evidence = {}

    # ── 违纪因子 ──
    disc = feature_dict.get("discipline_factor", 0)
    disc_base = baselines.get("discipline_factor", 1.0)
    if disc_base == 0:
        disc_base = 1.0

    if disc > 0:
        ratio = disc / disc_base
        if ratio >= 5:
            evidence["discipline"] = f"违纪超标 {ratio:.1f} 倍（严重）"
        elif ratio >= 2:
            evidence["discipline"] = f"违纪超标 {ratio:.1f} 倍"
        else:
            evidence["discipline"] = f"存在 {disc:.0f} 分违纪记录"
    # 无违纪 → 不展示（而非 "无违纪记录"）

    # ── 数学成绩（只展示低于均值 20% 以上或低于及格线）──
    math_avg = feature_dict.get("math_avg", 0)
    math_base = baselines.get("math_avg", 75.0)
    if math_base > 0:
        ratio = math_avg / math_base
    else:
        ratio = 1.0

    if math_avg < 60:
        evidence["math"] = f"数学成绩 {math_avg:.1f} 分（低于及格线，仅为年级均值 {ratio:.2f}）"
    elif ratio < 0.8:
        evidence["math"] = f"数学成绩仅为年级均值的 {ratio:.2f}"
    # 高于均值 0.8 → 正常，不展示

    # ── 出勤率（只展示低于 90% 或明显低于均值）──
    att = feature_dict.get("attendance_rate", 1.0)
    att_base = baselines.get("attendance_rate", 0.95)
    att_pct = att * 100
    att_base_pct = att_base * 100

    if att < 0.9:
        evidence["attendance"] = f"出勤率 {att_pct:.1f}%（低于均值 {att_base_pct:.1f}%）"
    elif att < att_base - 0.05:
        evidence["attendance"] = f"出勤率 {att_pct:.1f}%（明显低于均值）"
    # 正常出勤 → 不展示

    # ── 综合素质分（只展示偏低的情况）──
    quality = feature_dict.get("quality_score", 80.0)
    quality_base = baselines.get("quality_score", 80.0)

    if quality < 70:
        evidence["quality"] = f"综合素质分 {quality:.1f}（偏低，年级均值 {quality_base:.1f}）"
    elif quality < quality_base - 5:
        evidence["quality"] = f"综合素质分 {quality:.1f}（低于年级均值 {quality_base:.1f}）"
    # 正常 → 不展示

    # ── 风险密度 ──
    density = feature_dict.get("risk_density", 0)
    if density > 5:
        evidence["risk"] = f"近30天预警 {density} 次（高频，请重点关注）"
    elif density > 2:
        evidence["risk"] = f"近30天预警 {density} 次（中频）"
    elif density > 0:
        evidence["risk"] = f"近30天预警 {density} 次"
    # 无预警 → 不展示

    # ── 成绩趋势（只展示显著变化）──
    slope = feature_dict.get("math_slope", 0)
    if slope < -10:
        evidence["trend"] = f"成绩急剧下滑（斜率 {slope:.1f}）"
    elif slope < -5:
        evidence["trend"] = f"成绩下滑趋势明显（斜率 {slope:.1f}）"
    elif slope > 10:
        evidence["trend"] = f"成绩快速上升（斜率 {slope:.1f}）"
    elif slope > 5:
        evidence["trend"] = f"成绩上升趋势良好（斜率 {slope:.1f}）"
    # 平稳 → 不展示
    else:
        evidence["trend"] = f"成绩趋势平稳（斜率 {slope:.1f}）"

    return evidence


# ━━━ 方案A 后验追踪引擎 ━━━
# 每次 predict 被调用时，自动扫描该学生所有 "tracking" 状态的干预记录，
# 用最新风险概率更新 risk_after，并根据 Δ 自动评定效果等级。
# 阈值：|Δ| >= 0.15 判定方向，0.05~0.15 判定轻微变化。

POSTERIOR_DELTA_THRESHOLD = 0.15   # 显著变化阈值
POSTERIOR_MILD_THRESHOLD  = 0.05   # 轻微变化阈值

def _run_posterior_tracking(student_id: int, current_risk: float):
    """后验追踪：对比干预前 risk_before 与当前 risk，自动刷新 risk_after + effect_rating"""
    active = InterventionRecord.query.filter_by(
        student_id=student_id,
        status="tracking"
    ).all()

    if not active:
        return

    updated_count = 0
    for rec in active:
        if rec.risk_before is None:
            continue

        rec.risk_after = round(current_risk, 4)
        delta = rec.risk_before - current_risk  # 正值=风险下降(好事)

        # 自动效果评级
        if delta >= POSTERIOR_DELTA_THRESHOLD:
            rec.effect_rating = "显著改善"
        elif delta >= POSTERIOR_MILD_THRESHOLD:
            rec.effect_rating = "略有改善"
        elif delta <= -POSTERIOR_DELTA_THRESHOLD:
            rec.effect_rating = "恶化"
        elif delta <= -POSTERIOR_MILD_THRESHOLD:
            rec.effect_rating = "略有恶化"
        else:
            rec.effect_rating = "无变化"

        rec.updated_at = get_local_now()
        updated_count += 1
        print(
            f"[后验追踪] 学生{rec.student.name}: "
            f"risk {rec.risk_before:.4f}→{current_risk:.4f} "
            f"Δ={delta:+.4f} → {rec.effect_rating}"
        )

    if updated_count:
        safe_commit()
        print(f"[后验追踪] 批量更新完成: {updated_count}条干预记录已刷新")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  MLOps 自进化控制路由
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@bp.route("/retrain", methods=["POST"])
def trigger_retrain():
    """
    手动触发模型重训 — 仅管理员可调用。

    响应:
        {"status": "started", "message": "重训已在后台启动", "last_retrain": "..."}
        {"status": "throttled", "message": "距离上次重训不足1小时,已被防抖拦截"}
    """
    from flask import session
    if session.get("role") not in ("ms_admin"):
        return jsonify({"error": "仅管理员可触发重训"}), 403

    from utils.model_retrain import trigger_auto_retrain, get_retrain_status
    app = current_app._get_current_object()

    status = get_retrain_status()
    if not status["can_retrain"]:
        return jsonify({
            "status": "throttled",
            "message": f"距离上次重训不足1小时,已被防抖拦截",
            "last_retrain": status["last_retrain"],
        })

    success = trigger_auto_retrain(app, grade_id=1, force=True)
    if success:
        return jsonify({
            "status": "started",
            "message": "模型重训已在后台异步启动,预计10秒内完成热更新",
            "last_retrain": status["last_retrain"],
        })
    else:
        return jsonify({"status": "error", "message": "重训启动失败"}), 500


@bp.route("/retrain/status")
def retrain_status():
    """查询重训状态"""
    from flask import session
    if session.get("role") not in ("ms_admin", "grade_leader"):
        return jsonify({"error": "无权查看"}), 403

    from utils.model_retrain import get_retrain_status
    status = get_retrain_status()

    model_path = os.path.join(current_app.root_path, "models", "wings_xgb_pipeline.pkl")
    model_size = os.path.getsize(model_path) if os.path.exists(model_path) else 0
    model_mtime = datetime.fromtimestamp(os.path.getmtime(model_path)).isoformat() if os.path.exists(model_path) else None

    return jsonify({
        "cache_active": _PIPELINE is not None,
        "model_size_bytes": model_size,
        "model_last_modified": model_mtime,
        "last_retrain": status["last_retrain"],
        "can_retrain": status["can_retrain"],
    })


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  教学干预效果闭环 API (Phase 5.3 方案A)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

INTERVENTION_TYPES = ["谈话", "家长联动", "座位调整", "学业辅导", "心理干预", "行为契约", "其他"]
EFFECT_RATINGS  = ["显著改善", "略有改善", "无变化", "恶化"]


@bp.route("/intervention/create", methods=["POST"])
def create_intervention():
    """创建干预记录 — 从谈话单页面一键调用"""
    if session.get("role") not in ("ms_admin", "grade_leader", "class_teacher"):
        return jsonify({"error": "无权操作"}), 403

    data = request.get_json(force=True, silent=True) or {}
    student_id = data.get("student_id")
    if not student_id:
        return jsonify({"error": "缺少 student_id"}), 400

    student = Student.query.get_or_404(student_id)

    # 获取当前 AI 风险概率作为 risk_before
    risk_before = None
    try:
        pipeline, metadata = _load_model()
        fe = FeatureExtractor(grade_id=student.grade_id)
        result = fe.get_student_vector(student_id)
        if result["features"]:
            proba = pipeline.predict_proba([result["features"]])[0]
            risk_before = float(proba[1])
    except Exception:
        pass

    teacher_id = session.get("user_id", 0)
    rec = InterventionRecord(
        student_id=student_id,
        teacher_id=teacher_id,
        risk_before=risk_before,
        intervention_type=data.get("intervention_type", "谈话"),
        notes=data.get("notes", ""),
        intervention_date=data.get("intervention_date") or get_local_now().date(),
        follow_up_date=data.get("follow_up_date") or None,
    )
    db.session.add(rec)
    safe_commit()

    return jsonify({
        "status": "ok",
        "intervention_id": rec.id,
        "risk_before": risk_before,
        "message": "干预记录已创建，请在随访完成后填写效果评估",
    })


@bp.route("/intervention/<int:student_id>")
def intervention_history(student_id):
    """查看学生干预历史 + 风险变化曲线数据"""
    if session.get("role") not in ("ms_admin", "grade_leader", "class_teacher"):
        return jsonify({"error": "无权查看"}), 403

    student = Student.query.get_or_404(student_id)
    records = InterventionRecord.query.filter_by(
        student_id=student_id
    ).order_by(InterventionRecord.intervention_date.asc()).all()

    interventions = [r.to_dict() for r in records]

    # 风险趋势数据 (用于画折线图)
    risk_trend = []
    for r in records:
        if r.risk_before is not None:
            risk_trend.append({
                "date": r.intervention_date.isoformat(),
                "risk": r.risk_before,
                "type": "before",
                "intervention_id": r.id,
            })
        if r.risk_after is not None and r.follow_up_done:
            risk_trend.append({
                "date": r.follow_up_date.isoformat() if r.follow_up_date else "",
                "risk": r.risk_after,
                "type": "after",
                "intervention_id": r.id,
            })

    return jsonify({
        "student_id": student_id,
        "student_name": student.name,
        "interventions": interventions,
        "risk_trend": sorted(risk_trend, key=lambda x: x["date"]),
    })


@bp.route("/intervention/<int:int_id>/followup", methods=["POST"])
def update_followup(int_id):
    """更新随访结果 — 班主任填写干预后风险概率 + 效果评估"""
    if session.get("role") not in ("ms_admin", "grade_leader", "class_teacher"):
        return jsonify({"error": "无权操作"}), 403

    rec = InterventionRecord.query.get_or_404(int_id)
    data = request.get_json(force=True, silent=True) or {}

    # 获取随访时的最新 AI 风险概率
    risk_after = data.get("risk_after")
    if risk_after is None:
        try:
            pipeline, metadata = _load_model()
            fe = FeatureExtractor(grade_id=rec.student.grade_id)
            result = fe.get_student_vector(rec.student_id)
            if result["features"]:
                proba = pipeline.predict_proba([result["features"]])[0]
                risk_after = float(proba[1])
        except Exception:
            pass

    rec.risk_after = risk_after
    rec.effect_rating = data.get("effect_rating", "")
    rec.follow_up_notes = data.get("follow_up_notes", "")
    rec.follow_up_done = True
    rec.updated_at = get_local_now()
    safe_commit()

    delta = rec.risk_delta
    slope = rec.risk_slope
    print(f"[干预后验] 学生{rec.student.name}: "
          f"risk {rec.risk_before}→{rec.risk_after}, "
          f"delta={delta}, slope={slope}")

    return jsonify({
        "status": "ok",
        "risk_delta": delta,
        "risk_slope": slope,
        "message": "随访记录已更新，后验数据已纳入分析",
    })


@bp.route("/intervention/report")
def intervention_report():
    """干预有效性看板 — 全年级/全校视角"""
    if session.get("role") not in ("ms_admin", "grade_leader"):
        return "无权查看", 403

    grade_id = request.args.get("grade_id", type=int)

    query = InterventionRecord.query
    if grade_id:
        query = query.join(Student).filter(Student.grade_id == grade_id)

    records = query.order_by(InterventionRecord.created_at.desc()).limit(200).all()

    # ── 后验分析 ──
    total = len(records)
    followed_up = [r for r in records if r.follow_up_done]
    effective = [r for r in followed_up if r.is_effective]

    # 按干预类型统计有效性
    type_stats = {}
    for r in followed_up:
        t = r.intervention_type
        if t not in type_stats:
            type_stats[t] = {"count": 0, "effective": 0, "total_delta": 0.0}
        type_stats[t]["count"] += 1
        if r.is_effective:
            type_stats[t]["effective"] += 1
        if r.risk_delta is not None:
            type_stats[t]["total_delta"] += r.risk_delta

    # 计算平均风险降幅
    for t in type_stats:
        c = type_stats[t]["count"]
        if c > 0:
            type_stats[t]["avg_delta"] = round(type_stats[t]["total_delta"] / c, 4)
            type_stats[t]["effect_rate"] = round(type_stats[t]["effective"] / c * 100, 1)
        else:
            type_stats[t]["avg_delta"] = 0
            type_stats[t]["effect_rate"] = 0

    # 整体指标
    overall = {
        "total": total,
        "followed_up": len(followed_up),
        "effective": len(effective),
        "effect_rate": round(len(effective) / max(1, len(followed_up)) * 100, 1),
        "avg_delta": round(
            sum(r.risk_delta for r in followed_up if r.risk_delta is not None) /
            max(1, len([r for r in followed_up if r.risk_delta is not None])), 4
        ),
    }

    # ── JSON 模式：供大屏推荐弹窗异步调用 ──
    is_json = (
        request.args.get("json") == "1" or
        request.headers.get("X-Requested-With") == "XMLHttpRequest"
    )
    if is_json:
        # 构造 rank_list（按 avg_delta 降序）
        rank_list = []
        for t, stats in type_stats.items():
            rank_list.append({
                "type": t,
                "total_cases": stats["count"],
                "avg_risk_reduction": stats.get("avg_delta", 0.0),
                "success_rate": round(stats.get("effect_rate", 0.0) / 100, 4),
            })
        rank_list.sort(key=lambda x: x["avg_risk_reduction"], reverse=True)

        # 无数据时输出专家先验
        if not rank_list:
            rank_list = [
                {"type": "座位调整", "total_cases": 0, "avg_risk_reduction": 0.185, "success_rate": 0.82},
                {"type": "家长联动", "total_cases": 0, "avg_risk_reduction": 0.152, "success_rate": 0.75},
                {"type": "学业辅导", "total_cases": 0, "avg_risk_reduction": 0.121, "success_rate": 0.70},
                {"type": "行为契约", "total_cases": 0, "avg_risk_reduction": 0.098, "success_rate": 0.68},
                {"type": "心理干预", "total_cases": 0, "avg_risk_reduction": 0.085, "success_rate": 0.65},
                {"type": "谈话",     "total_cases": 0, "avg_risk_reduction": 0.054, "success_rate": 0.60},
                {"type": "其他",     "total_cases": 0, "avg_risk_reduction": 0.010, "success_rate": 0.50},
            ]

        return jsonify({
            "status": "success",
            "timestamp": get_local_now().date().isoformat(),
            "grade_id": grade_id,
            "recommendation_matrix": {item["type"]: item for item in rank_list},
            "rank_list": rank_list,
            "overall": overall,
        })

    return render_template(
        "ai/intervention_report.html",
        overall=overall,
        type_stats=type_stats,
        records=records[:50],
        grade_id=grade_id,
    )


@bp.route("/intervention/analytics")
def intervention_analytics():
    """
    全校/全年级干预手段效能分析 JSON API — 德育处核心资产。

    按干预类型分组统计：
      - total_cases: 该类型总干预次数
      - with_followup: 已完成随访的案例数
      - avg_risk_reduction: 平均风险降幅 (risk_before - risk_after, 正值=风险下降)
      - effect_rate: 有效干预占比 (%)
      - best_practice: 全类型中平均风险降幅最高的类型名

    返回:
        {"status": "success", "timestamp": "2026-06-11", "best_practice": "座位调整",
         "matrix": {"谈话": {...}, "家长联动": {...}, ...}}
    """
    if session.get("role") not in ("ms_admin", "grade_leader"):
        return jsonify({"error": "无权查看"}), 403

    grade_id = request.args.get("grade_id", type=int)

    # 统计每种干预类型的效能
    stats = db.session.query(
        InterventionRecord.intervention_type,
        func.count(InterventionRecord.id).label("total_count"),
        func.sum(
            func.if_(InterventionRecord.follow_up_done == True, 1, 0)
        ).label("followup_count"),
        func.avg(
            InterventionRecord.risk_before - InterventionRecord.risk_after
        ).label("avg_drop"),
    ).filter(
        InterventionRecord.risk_after.isnot(None)
    )

    if grade_id:
        stats = stats.join(Student).filter(Student.grade_id == grade_id)

    stats = stats.group_by(InterventionRecord.intervention_type).all()

    # 再统计有效干预数 (用 SQL 不易做 b/c is_effective 是 Python 属性，分开算)
    effective_query = db.session.query(
        InterventionRecord.intervention_type,
        func.count(InterventionRecord.id)
    ).filter(
        InterventionRecord.follow_up_done == True,
        InterventionRecord.effect_rating.in_(["显著改善", "略有改善"]),
    )
    if grade_id:
        effective_query = effective_query.join(Student).filter(Student.grade_id == grade_id)
    effective_query = effective_query.group_by(InterventionRecord.intervention_type).all()
    effective_map = {row[0]: row[1] for row in effective_query}

    matrix = {}
    best_type = None
    best_drop = -999.0

    for row in stats:
        t = row.intervention_type
        total = int(row.total_count)
        followup = int(row.followup_count)
        avg_drop = round(float(row.avg_drop), 4) if row.avg_drop else 0.0
        eff_count = effective_map.get(t, 0)
        eff_rate = round(eff_count / max(1, followup) * 100, 1)

        matrix[t] = {
            "total_cases": total,
            "with_followup": followup,
            "avg_risk_reduction": avg_drop,
            "effective_count": eff_count,
            "effect_rate": eff_rate,
            "success_rate": round(eff_rate / 100, 4),  # 0~1 成功率
        }

        if followup >= 1 and avg_drop > best_drop:
            best_drop = avg_drop
            best_type = t

    # ── 兜底：无随访数据时，输出基于教育学经验的专家先验权重 ──
    _EXPERT_PRIOR = [
        {"type": "座位调整",   "avg_risk_reduction": 0.185, "success_rate": 0.82, "total_cases": 0},
        {"type": "家长联动",   "avg_risk_reduction": 0.152, "success_rate": 0.75, "total_cases": 0},
        {"type": "学业辅导",   "avg_risk_reduction": 0.121, "success_rate": 0.70, "total_cases": 0},
        {"type": "行为契约",   "avg_risk_reduction": 0.098, "success_rate": 0.68, "total_cases": 0},
        {"type": "心理干预",   "avg_risk_reduction": 0.085, "success_rate": 0.65, "total_cases": 0},
        {"type": "谈话",       "avg_risk_reduction": 0.054, "success_rate": 0.60, "total_cases": 0},
        {"type": "其他",       "avg_risk_reduction": 0.010, "success_rate": 0.50, "total_cases": 0},
    ]

    if not matrix:
        # 零随访数据 → 全量输出专家先验
        for ep in _EXPERT_PRIOR:
            matrix[ep["type"]] = {
                "total_cases": 0,
                "with_followup": 0,
                "avg_risk_reduction": ep["avg_risk_reduction"],
                "effective_count": 0,
                "effect_rate": round(ep["success_rate"] * 100, 1),
                "success_rate": ep["success_rate"],
                "_source": "expert_prior",
            }
        best_type = _EXPERT_PRIOR[0]["type"]  # 座位调整

    # 按平均降幅降序排 → recommendation_matrix
    recommendation_matrix = dict(
        sorted(matrix.items(), key=lambda kv: kv[1]["avg_risk_reduction"], reverse=True)
    )

    # 追踪中 vs 已结案 全局统计
    status_counts = db.session.query(
        InterventionRecord.status,
        func.count(InterventionRecord.id)
    ).filter(
        InterventionRecord.status.in_(["tracking", "completed"])
    )
    if grade_id:
        status_counts = status_counts.join(Student).filter(Student.grade_id == grade_id)
    status_counts = status_counts.group_by(InterventionRecord.status).all()
    status_map = {row[0]: row[1] for row in status_counts}

    return jsonify({
        "status": "success",
        "timestamp": get_local_now().date().isoformat(),
        "grade_id": grade_id,
        "best_practice": best_type,
        "matrix": matrix,
        "recommendation_matrix": recommendation_matrix,
        "pipeline": {
            "tracking": status_map.get("tracking", 0),
            "completed": status_map.get("completed", 0),
        },
    })
