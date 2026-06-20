# -*- coding: utf-8 -*-
"""
AI 评语引擎 — 统一蓝图（合并 endterm_comment + value_added）

功能矩阵：
  [期末评语]  撰写/批量录入/查看/导出 + AI生成舱
  [增值评价]  隐形好学生扫描/AI温暖评语/闪光标签

架构：
  - 底层数据聚合 → utils/student_data_aggregator.py（共享）
  - LLM 调用      → utils/llm_client（统一熔断）
  - Prompt 切换   → style_flag: formal / warm
"""
from flask import (
    Blueprint, render_template, request, redirect, url_for,
    flash, session, jsonify, current_app, make_response,
)
from models import db, Student, Class, Grade, EndTermComment, ValueAddedComment, User, Exam
from decorators import login_required, require_role
from utils import get_local_now
from utils.db_utils import safe_commit
from utils.llm_client import call_llm_json, LLMAvailabilityError
from utils.student_data_aggregator import StudentDataAggregator as SDA
from blueprints.audit_log import audit_log
import json
import time


ai_comment_bp = Blueprint("ai_comment", __name__)

# ══════════════════════════════════════════════════════════════
#  System Prompts — 风格切换矩阵
# ══════════════════════════════════════════════════════════════

FORMAL_SYSTEM_PROMPT = """你是一位拥有30年经验的初中德育与学情因果诊断专家。你的任务是根据学生的数字特征上下文（包含大考斜率、5维行为分、违纪记录、活动参与和班主任随访文本），输出高度定制化、具备绝对人文关怀与循证支撑的期末评语。

[核心约束]
1. 严禁出现"该生团结同学、热爱劳动、尊敬师长"等万能套话。
2. 必须进行【跨界交叉论证】：如果成绩斜率向上且行为分中"智"维度高，必须点出其考前付出的努力；如果行为高频微量下滑，评语中必须隐晦且温柔地给予心理警示。
3. 语气必须是"既有老教师的慈爱和深刻，又有数据科学的精准"。
4. 评语中必须包含至少1个具体数据点（如"数学得分率向上拉升了X个百分点"）和至少1个具体行为事件（如"主动擦黑板"）。
5. overall_comment 字数严格控制在180-280字之间，既要有温度又要有信息密度。

[输出格式]
必须输出严格的 JSON 格式，不要任何 Markdown 包裹（如 ```json），直接返回以下结构的字符串：
{
  "strengths": "一句话精准提炼核心闪光点（结合真实事件或数据，15-30字）",
  "improvements": "一句话指出下一步需要突破的盲区（如特定知识点或行为习惯，15-30字）",
  "overall_comment": "180-280字的史诗级个性化评语，用于打印落盘",
  "teacher_suggestion": "给班主任的私密带班建议（如：该生自尊心极强，不宜公开批评，建议采用空间隔离或导师一帮一）"
}"""

WARM_SYSTEM_PROMPT = """你是一位拥有25年经验的初中德育主任兼心理咨询师，专门关注那些"隐形好学生"——成绩不在前列但品行端正、正在默默进步的孩子。你的任务是根据学生的数据画像，生成一段温暖、有力、让人落泪的鼓励评语。

[核心约束]
1. 严禁出现"该生学习态度端正"等万能套话。每句话必须有数据或事实支撑。
2. 必须用"你"而非"该生"，语气像一位慈爱而睿智的长辈直接对话。
3. 必须点出至少1个具体进步（如"数学从X分提到了Y分"）和1个品行亮点。
4. 结尾必须有一句温暖展望，让学生感受到被看见、被认可。
5. comment_text 字数严格控制在 150-250 字之间。
6. 评语要让阅读的学生感到"原来老师一直在看着我"。

[输出格式]
必须输出严格的 JSON 格式，不要任何 Markdown 包裹（如 ```json），直接返回以下结构的字符串：
{
  "comment_text": "150-250字的温暖鼓励评语",
  "highlight_tags": ["闪光点标签1", "闪光点标签2", "闪光点标签3"]
}"""

PROMPT_MAP = {
    "formal": FORMAL_SYSTEM_PROMPT,
    "warm": WARM_SYSTEM_PROMPT,
}


# ══════════════════════════════════════════════════════════════
#  统一首页 — 展示两种评语类型
# ══════════════════════════════════════════════════════════════

@ai_comment_bp.route("/")
@login_required
def index():
    """AI评语统一首页"""
    return render_template("ai_comment/index.html")


# ══════════════════════════════════════════════════════════════
#  期末评语子模块
# ══════════════════════════════════════════════════════════════

@ai_comment_bp.route("/endterm")
@login_required
def endterm_list():
    class_id = session.get("class_id")
    grade_id = session.get("grade_id")
    role = session.get("role")

    filter_class_id = request.args.get("class_id", "", type=int) or ""
    status_filter = request.args.get("status", "")
    semester_filter = request.args.get("semester", "")

    semesters = sorted(
        {r[0] for r in db.session.query(EndTermComment.semester).distinct().all()},
        reverse=True,
    )
    current_semester = semester_filter or (semesters[0] if semesters else "2025-2026-2")

    if role == "ms_admin":
        classes = Class.query.filter_by(is_active=True).order_by(Class.name).all()
    elif role == "grade_leader":
        classes = Class.query.filter_by(grade_id=grade_id, is_active=True).order_by(Class.name).all()
    else:
        classes = Class.query.filter_by(id=class_id, is_active=True).order_by(Class.name).all()

    query = EndTermComment.query.join(Student)
    if role == "class_teacher":
        query = query.filter(Student.class_id == class_id)
    elif role == "grade_leader":
        query = query.filter(Student.grade_id == grade_id)
    if current_semester:
        query = query.filter(EndTermComment.semester == current_semester)
    if filter_class_id:
        query = query.filter(Student.class_id == filter_class_id)
    if status_filter:
        query = query.filter(EndTermComment.status == status_filter)

    comments = query.order_by(EndTermComment.created_at.desc()).all()
    return render_template(
        "ai_comment/endterm_list.html",
        comments=comments,
        semesters=semesters,
        current_semester=current_semester,
        classes=classes,
        filter_class_id=filter_class_id,
        status_filter=status_filter,
    )


@ai_comment_bp.route("/endterm/create", methods=["GET", "POST"])
@login_required
@require_role("class_teacher")
@audit_log("create_comment", "EndTermComment")
def endterm_create():
    class_id = session.get("class_id")
    students = []
    if class_id:
        students = Student.query.filter_by(
            class_id=class_id, is_active=True
        ).order_by(Student.student_no).all()

    if request.method == "POST":
        student_id = request.form.get("student_id", "")
        overall_comment = request.form.get("overall_comment", "")
        if not student_id or not overall_comment:
            flash("请选择学生并填写综合评语", "danger")
            return redirect(url_for("ai_comment.endterm_create"))

        student = Student.query.get(int(student_id))
        if not student:
            flash("学生不存在", "danger")
            return redirect(url_for("ai_comment.endterm_create"))

        c = EndTermComment(
            student_id=int(student_id),
            class_id=student.class_id,
            grade_id=student.grade_id,
            overall_comment=overall_comment,
            strengths=request.form.get("strengths", ""),
            improvements=request.form.get("improvements", ""),
            teacher_suggestion=request.form.get("teacher_suggestion", ""),
            semester=request.form.get("semester", SDA.current_semester()),
            created_by=session.get("display_name") or session.get("username", ""),
        )
        db.session.add(c)
        safe_commit()
        flash("评语已保存", "success")
        return redirect(url_for("ai_comment.endterm_list"))

    pre_student_id = request.args.get("student_id", type=int)
    student_context = None
    student = None
    if pre_student_id:
        student = Student.query.get(pre_student_id)
        if student:
            student_context = SDA.build_student_context(student)

    return render_template(
        "ai_comment/endterm_form.html",
        students=students,
        comment=None,
        student=student,
        student_context=student_context,
        pre_student_id=pre_student_id or "",
        semester=SDA.current_semester(),
    )


@ai_comment_bp.route("/endterm/<int:cid>/edit", methods=["GET", "POST"])
@login_required
def endterm_edit(cid):
    c = EndTermComment.query.get_or_404(cid)
    if session.get("role") == "class_teacher" and c.class_id != session.get("class_id"):
        flash("无权操作", "danger")
        return redirect(url_for("ai_comment.endterm_list"))

    if request.method == "POST":
        c.overall_comment = request.form.get("content", c.overall_comment)
        c.strengths = request.form.get("strengths", c.strengths)
        c.improvements = request.form.get("improvements", c.improvements)
        c.teacher_suggestion = request.form.get("teacher_suggestion", c.teacher_suggestion)
        c.status = "published" if request.form.get("is_published") == "on" else "draft"
        safe_commit()
        flash("已更新", "success")
        return redirect(url_for("ai_comment.endterm_list"))

    class_id = session.get("class_id")
    students = Student.query.filter_by(
        class_id=class_id, is_active=True
    ).order_by(Student.student_no).all()

    student_context = None
    student = c.student if c else None
    if student:
        student_context = SDA.build_student_context(student)

    return render_template(
        "ai_comment/endterm_form.html",
        students=students,
        comment=c,
        student=student,
        student_context=student_context,
        pre_student_id=c.student_id,
        semester=c.semester,
    )


@ai_comment_bp.route("/endterm/<int:cid>/delete", methods=["POST"])
@login_required
def endterm_delete(cid):
    c = EndTermComment.query.get_or_404(cid)
    if session.get("role") == "class_teacher" and c.class_id != session.get("class_id"):
        flash("无权操作", "danger")
        return redirect(url_for("ai_comment.endterm_list"))
    db.session.delete(c)
    safe_commit()
    flash("已删除", "info")
    return redirect(url_for("ai_comment.endterm_list"))


@ai_comment_bp.route("/endterm/batch", methods=["GET", "POST"])
@login_required
@require_role("class_teacher")
def endterm_batch():
    if request.method == "POST":
        return jsonify({"ok": True})
    return render_template("ai_comment/endterm_batch.html")


@ai_comment_bp.route("/endterm/export")
@login_required
def endterm_export():
    semester = request.args.get("semester", "")
    class_id = request.args.get("class_id", type=int)
    type_ = request.args.get("type", "word")

    q = EndTermComment.query
    if session.get("role") == "class_teacher":
        class_id = session.get("class_id")
    if class_id:
        q = q.join(Student).filter(Student.class_id == class_id)
    comments = q.order_by(
        EndTermComment.semester.desc(), EndTermComment.id.asc()
    ).all()

    if not comments:
        flash("暂无评语数据可导出", "warning")
        return redirect(url_for("ai_comment.endterm_list"))

    if type_ == "word":
        return _export_word(comments, semester)
    else:
        return _export_html(comments, semester)


def _export_word(comments, semester):
    """生成 .docx 导出"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        flash("服务器未安装 python-docx，请联系管理员安装", "danger")
        return redirect(url_for("ai_comment.endterm_list"))

    doc = Document()
    h = doc.add_heading("梨江中学 期末评语", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"学期：{semester or '全部'}")
    doc.add_paragraph("")

    for c in comments:
        s = c.student
        p = doc.add_paragraph()
        p.add_run(f"【{s.name if s else '?'}】  ").bold = True
        if s and s.class_:
            p.add_run(f"{s.class_.name}  ")
        doc.add_paragraph(f"综合评语：{c.overall_comment or '—'}")
        if c.strengths:
            doc.add_paragraph(f"主要优点：{c.strengths}")
        if c.improvements:
            doc.add_paragraph(f"改进建议：{c.improvements}")
        doc.add_paragraph("—" * 30)

    import io
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    resp = make_response(buf.read())
    resp.headers["Content-Type"] = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    resp.headers["Content-Disposition"] = (
        f"attachment; filename=期末评语_{semester or '全部'}.docx"
    )
    return resp


def _export_html(comments, semester):
    """HTML 打印版 — 可 Ctrl+P 保存 PDF"""
    html = render_template(
        "ai_comment/endterm_print.html", comments=comments, semester=semester
    )
    resp = make_response(html)
    resp.headers["Content-Type"] = "text/html; charset=utf-8"
    return resp


# ══════════════════════════════════════════════════════════════
#  AI 评语生成 — 统一端点（style_flag 切换 Prompt）
# ══════════════════════════════════════════════════════════════

@ai_comment_bp.route("/api/generate/<int:student_id>", methods=["POST"])
@login_required
@require_role("class_teacher", "ms_admin", "grade_leader")
def ai_generate(student_id):
    """统一 AI 评语生成端点

    Query params:
        style:  "formal" (期末正式评语, 4字段) | "warm" (温暖鼓励评语, 2字段)
        默认: "formal"
    """
    stu = Student.query.get_or_404(student_id)

    # 班主任数据隔离
    user_role = session.get("role")
    if user_role == "class_teacher":
        teacher_class = session.get("class_id")
        if teacher_class and stu.class_id != teacher_class:
            return jsonify({"status": "error", "message": "无权访问该学生"}), 403

    style = request.args.get("style", "formal")
    system_prompt = PROMPT_MAP.get(style, FORMAL_SYSTEM_PROMPT)

    # 共享数据聚合器
    context = SDA.build_llm_context(stu)
    user_prompt = SDA.format_llm_prompt(stu, context)

    try:
        if style == "warm":
            result = call_llm_json(system_prompt, user_prompt, max_tokens=1024, timeout=45)
        else:
            result = call_llm_json(system_prompt, user_prompt, max_tokens=1024, timeout=30)

        if style == "warm":
            if "comment_text" not in result:
                result["comment_text"] = ""
            if "highlight_tags" not in result:
                result["highlight_tags"] = []
        else:
            for key in ["strengths", "improvements", "overall_comment", "teacher_suggestion"]:
                if key not in result:
                    result[key] = ""

        return jsonify({
            "status": "success",
            "style": style,
            "data": result,
            "student_name": stu.name,
            "student_id": stu.id,
        })

    except LLMAvailabilityError as e:
        return jsonify({"status": "error", "message": str(e)}), 503
    except (ValueError, json.JSONDecodeError) as e:
        current_app.logger.error(f"LLM返回非JSON格式: {str(e)[:200]}")
        return jsonify({"status": "error", "message": "大模型返回格式异常，请重试"}), 500
    except Exception as e:
        current_app.logger.error(f"评语生成未知错误: {str(e)}")
        return jsonify({"status": "error", "message": "评语织网机遭遇未知故障"}), 500


@ai_comment_bp.route("/api/batch-generate", methods=["POST"])
@login_required
@require_role("class_teacher")
def ai_batch_generate():
    """全班批量生成期末评语 (style=formal)"""
    class_id = session.get("class_id")
    if not class_id:
        return jsonify({"status": "error", "message": "无法获取您的班级"}), 400

    style = request.args.get("style", "formal")
    system_prompt = PROMPT_MAP.get(style, FORMAL_SYSTEM_PROMPT)

    students = Student.query.filter_by(
        class_id=class_id, is_active=True
    ).order_by(Student.student_no).all()

    if not students:
        return jsonify({"status": "error", "message": "班级无在读学生"}), 404

    results = []
    success_count = 0
    fail_count = 0

    for stu in students:
        try:
            context = SDA.build_llm_context(stu)
            user_prompt = SDA.format_llm_prompt(stu, context)
            comment_data = call_llm_json(system_prompt, user_prompt, max_tokens=1024, timeout=30)
            results.append({
                "student_id": stu.id,
                "student_name": stu.name,
                "status": "success",
                "data": comment_data,
            })
            success_count += 1
        except Exception as e:
            results.append({
                "student_id": stu.id,
                "student_name": stu.name,
                "status": "error",
                "message": str(e)[:100],
            })
            fail_count += 1

    return jsonify({
        "status": "success" if fail_count == 0 else "partial",
        "total": len(students),
        "success": success_count,
        "failed": fail_count,
        "results": results,
    })


# ══════════════════════════════════════════════════════════════
#  增值评价子模块
# ══════════════════════════════════════════════════════════════

@ai_comment_bp.route("/value-added")
@require_role("ms_admin")
def value_added_list():
    """增值评价列表"""
    grade_id = request.args.get("grade_id", type=int)
    exam_id = request.args.get("exam_id", type=int)

    grades = Grade.query.all()
    exams = Exam.query.order_by(Exam.exam_date.desc()).all()

    q = ValueAddedComment.query.order_by(ValueAddedComment.created_at.desc())
    if grade_id:
        q = q.filter_by(grade_id=grade_id)
    if exam_id:
        q = q.filter_by(exam_id=exam_id)
    comments = q.limit(100).all()

    for c in comments:
        try:
            c.tags = json.loads(c.highlight_tags) if c.highlight_tags else []
        except Exception:
            c.tags = []

    return render_template(
        "ai_comment/value_added_list.html",
        comments=comments,
        grades=grades,
        exams=exams,
        grade_filter=grade_id,
        exam_filter=exam_id,
    )


@ai_comment_bp.route("/value-added/scan", methods=["POST"])
@require_role("ms_admin")
def value_added_scan():
    """扫描隐形好学生"""
    prev_exam_id = request.form.get("prev_exam_id", type=int)
    curr_exam_id = request.form.get("curr_exam_id", type=int)
    grade_id = request.form.get("grade_id", type=int)

    if not prev_exam_id or not curr_exam_id:
        flash("请选择两次考试", "danger")
        return redirect(url_for("ai_comment.value_added_list"))

    gems = SDA.identify_hidden_gems(prev_exam_id, curr_exam_id, grade_id)
    exams = {e.id: e.name for e in Exam.query.all()}

    return render_template(
        "ai_comment/value_added_scan.html",
        gems=gems,
        prev_exam_id=prev_exam_id,
        curr_exam_id=curr_exam_id,
        grade_id=grade_id,
        prev_exam_name=exams.get(prev_exam_id, ""),
        curr_exam_name=exams.get(curr_exam_id, ""),
        grades=Grade.query.all(),
    )


@ai_comment_bp.route("/value-added/generate", methods=["POST"])
@require_role("ms_admin")
def value_added_generate():
    """批量生成 AI 温暖评语"""
    prev_exam_id = request.form.get("prev_exam_id", type=int)
    curr_exam_id = request.form.get("curr_exam_id", type=int)
    grade_id = request.form.get("grade_id", type=int)
    student_ids = request.form.getlist("student_ids")

    if not student_ids:
        flash("请至少选择一个学生", "warning")
        return redirect(
            url_for(
                "ai_comment.value_added_scan",
                prev_exam_id=prev_exam_id,
                curr_exam_id=curr_exam_id,
                grade_id=grade_id,
            )
        )

    gems = SDA.identify_hidden_gems(prev_exam_id, curr_exam_id, grade_id)
    gem_map = {g["student"].id: g for g in gems}

    generated = 0
    failed = 0

    for sid_str in student_ids:
        sid = int(sid_str)
        gem = gem_map.get(sid)
        if not gem:
            continue

        s = gem["student"]

        existing = ValueAddedComment.query.filter_by(
            student_id=sid, exam_id=curr_exam_id
        ).first()
        if existing:
            db.session.delete(existing)
            db.session.flush()

        try:
            context = SDA.build_gem_context(gem)
            data = call_llm_json(WARM_SYSTEM_PROMPT, context, timeout=45)

            vac = ValueAddedComment(
                student_id=sid,
                class_id=s.class_id,
                grade_id=s.grade_id,
                exam_id=curr_exam_id,
                prev_total=gem["prev_total"],
                curr_total=gem["curr_total"],
                score_delta=gem["score_delta"],
                rank_delta=gem["prev_rank"] - gem["curr_rank"],
                behavior_score=gem["behavior_score"],
                discipline_count=gem["discipline_count"],
                attendance_rate=gem["attendance_rate"],
                comment_text=data.get("comment_text", ""),
                highlight_tags=json.dumps(data.get("highlight_tags", []), ensure_ascii=False),
                status="draft",
            )
            db.session.add(vac)
            generated += 1
            time.sleep(0.5)

        except Exception as e:
            failed += 1
            flash(f"{s.name}: AI 生成失败 - {str(e)[:50]}", "danger")

    safe_commit()

    if generated > 0:
        audit_log("generate_value_added", f"生成 {generated} 条增值评价（考试ID={curr_exam_id}）")
        flash(
            f"成功生成 {generated} 条温暖评语"
            + (f"，{failed} 条失败" if failed else ""),
            "success",
        )
    else:
        flash("未能生成任何评语", "danger")

    return redirect(url_for("ai_comment.value_added_list", exam_id=curr_exam_id))


@ai_comment_bp.route("/value-added/<int:cid>/publish", methods=["POST"])
@require_role("ms_admin")
def value_added_publish(cid):
    vac = ValueAddedComment.query.get_or_404(cid)
    vac.status = "published"
    safe_commit()
    flash("评语已发布", "success")
    return redirect(url_for("ai_comment.value_added_list"))


@ai_comment_bp.route("/value-added/<int:cid>/delete", methods=["POST"])
@require_role("ms_admin")
def value_added_delete(cid):
    vac = ValueAddedComment.query.get_or_404(cid)
    db.session.delete(vac)
    safe_commit()
    flash("评语已删除", "success")
    return redirect(url_for("ai_comment.value_added_list"))
