"""期末评语 — 班主任撰写/批量录入/查看/导出 + AI评语生成舱"""
from flask import Blueprint, render_template, request, redirect, url_for, flash, session, jsonify, current_app
from models import db, Student, Class, Grade, EndTermComment, User, \
    DisciplineRecord, Score, Exam, ActivityRegistration, Activity, MentalHealthAssessment, \
    WingsScore, InterventionRecord, TeacherNote
from decorators import login_required, require_role
import time
from datetime import datetime
from utils import get_local_now
from utils.db_utils import safe_commit
from blueprints.audit_log import audit_log
from sqlalchemy import func
import json
import requests

endterm_comment_bp = Blueprint("endterm_comment", __name__)

# ══════════════════════════════════════════════════════════════
#  AI 评语生成舱 — System Prompt 工业级编织矩阵
# ══════════════════════════════════════════════════════════════
COMMENT_EXPERT_SYSTEM_PROMPT = """你是一位拥有30年经验的初中德育与学情因果诊断专家。你的任务是根据学生的数字特征上下文（包含大考斜率、5维行为分、违纪记录、活动参与和班主任随访文本），输出高度定制化、具备绝对人文关怀与循证支撑的期末评语。

[核心约束]
1. 严禁出现"该生团结同学、热爱劳动、尊敬师长"等万能套话。
2. 必须进行【跨界交叉论证】：如果成绩斜率向上且行为分中"智"维度高，必须点出其考前付出的努力；如果行为高频微量下滑，评语中必须隐晦且温柔地给予心理警示。
3. 语气必须是"既有老教师的慈爱和深刻，又有数据科学的精准"。
4. 评语中必须包含至少1个具体数据点（如"数学得分率向上拉升了X个百分点"）和至少1个具体行为事件（如"主动擦黑板"）。
5. overall_comment 字数严格控制在180-280字之间，既要有温度又要有信息密度。

[输出格式]
必须输出严格的 JSON 格式，不要任何 Markdown 包裹（如 ```json），直接返回以下结构的字符串：
{
  "strengths": "一句话精准提炼核心闪光点（结合真实事件或数据，15-30字）",
  "improvements": "一句话指出下一步需要突破的盲区（如特定知识点或行为习惯，15-30字）",
  "overall_comment": "180-280字的史诗级个性化评语，用于打印落盘",
  "teacher_suggestion": "给班主任的私密带班建议（如：该生自尊心极强，不宜公开批评，建议采用空间隔离或导师一帮一）"
}"""


def _call_llm_api(system_prompt, user_content):
    """调用大模型API（支持DeepSeek/通义千问/OpenAI兼容接口）"""
    api_key = current_app.config.get("LLM_API_KEY", "")
    api_url = current_app.config.get("LLM_API_URL", "https://api.deepseek.com/v1/chat/completions")
    model = current_app.config.get("LLM_MODEL", "deepseek-chat")
    timeout = current_app.config.get("LLM_TIMEOUT", 30)

    if not api_key:
        raise RuntimeError("LLM_API_KEY 未配置，请在环境变量中设置")

    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content}
        ],
        "temperature": 0.7,
        "max_tokens": 1024,
        "response_format": {"type": "json_object"}
    }
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    resp = requests.post(api_url, json=payload, headers=headers, timeout=timeout)
    resp.raise_for_status()
    result = resp.json()
    return result["choices"][0]["message"]["content"]


def _current_semester():
    now = get_local_now()
    y = now.year
    m = now.month
    if m >= 9:
        return f"{y}-{y+1}-1"
    elif m >= 2:
        return f"{y-1}-{y}-2"
    else:
        return f"{y-1}-{y}-2"


@endterm_comment_bp.route("/")
@login_required
def index():
    class_id = session.get("class_id")
    grade_id = session.get("grade_id")
    role = session.get("role")

    # 筛选参数
    filter_class_id = request.args.get("class_id", "", type=int) or ""
    status_filter = request.args.get("status", "")
    semester_filter = request.args.get("semester", "")

    # 可用学期列表
    semesters = sorted(
        {r[0] for r in db.session.query(EndTermComment.semester).distinct().all()},
        reverse=True
    )
    current_semester = semester_filter or (semesters[0] if semesters else "2025-2026-2")

    # 可用班级
    if role == "ms_admin":
        classes = Class.query.filter_by(is_active=True).order_by(Class.name).all()
    elif role == "grade_leader":
        classes = Class.query.filter_by(grade_id=grade_id, is_active=True).order_by(Class.name).all()
    else:
        classes = Class.query.filter_by(id=class_id, is_active=True).order_by(Class.name).all()

    # 构建查询
    query = EndTermComment.query.join(Student)
    if role == "class_teacher":
        query = query.filter(Student.class_id == class_id)
    elif role == "grade_leader":
        query = query.filter(Student.grade_id == grade_id)
    if current_semester:
        query = query.filter(EndTermComment.semester == current_semester)
    if filter_class_id:
        query = query.filter(Student.class_id == filter_class_id)
    if status_filter:
        query = query.filter(EndTermComment.status == status_filter)

    comments = query.order_by(EndTermComment.created_at.desc()).all()
    return render_template("endterm_comment/index.html",
        comments=comments, semesters=semesters, current_semester=current_semester,
        classes=classes, filter_class_id=filter_class_id, status_filter=status_filter)


@endterm_comment_bp.route("/create", methods=["GET", "POST"])
@login_required
@require_role("class_teacher")
@audit_log("create_comment", "EndTermComment")
def create():
    class_id = session.get("class_id")
    students = []
    if class_id:
        students = Student.query.filter_by(class_id=class_id, is_active=True).order_by(
            Student.student_no
        ).all()

    if request.method == "POST":
        student_id = request.form.get("student_id", "")
        overall_comment = request.form.get("overall_comment", "")
        strengths = request.form.get("strengths", "")
        improvements = request.form.get("improvements", "")
        teacher_suggestion = request.form.get("teacher_suggestion", "")
        if not student_id or not overall_comment:
            flash("请选择学生并填写综合评语", "danger")
            return redirect(url_for("endterm_comment.create"))
        student = Student.query.get(int(student_id))
        if not student:
            flash("学生不存在", "danger")
            return redirect(url_for("endterm_comment.create"))
        c = EndTermComment(
            student_id=int(student_id),
            class_id=student.class_id,
            grade_id=student.grade_id,
            overall_comment=overall_comment,
            strengths=strengths,
            improvements=improvements,
            teacher_suggestion=teacher_suggestion,
            semester=request.form.get("semester", _current_semester()),
            created_by=session.get("display_name") or session.get("username", ""),
        )
        db.session.add(c)
        safe_commit()
        flash("评语已保存", "success")
        return redirect(url_for("endterm_comment.index"))

    # GET: 如果 URL 中指定了 student_id，拉取该学生的多维度数据
    pre_student_id = request.args.get("student_id", type=int)
    student_context = None
    student = None
    if pre_student_id:
        student = Student.query.get(pre_student_id)
        if student:
            student_context = _build_student_context(student)

    return render_template("endterm_comment/form.html",
                           students=students,
                           comment=None,
                           student=student,
                           student_context=student_context,
                           pre_student_id=pre_student_id or "",
                           semester=_current_semester())


@endterm_comment_bp.route("/<int:cid>/edit", methods=["GET", "POST"])
@login_required
def edit(cid):
    c = EndTermComment.query.get_or_404(cid)
    # 班主任只能编辑本班评语
    if session.get("role") == "class_teacher" and c.class_id != session.get("class_id"):
        flash("无权操作", "danger")
        return redirect(url_for("endterm_comment.index"))
    if request.method == "POST":
        c.overall_comment = request.form.get("content", c.overall_comment)
        c.status = "published" if request.form.get("is_published") == "on" else "draft"
        safe_commit()
        flash("已更新", "success")
        return redirect(url_for("endterm_comment.index"))

    class_id = session.get("class_id")
    students = Student.query.filter_by(class_id=class_id, is_active=True).order_by(
        Student.student_no
    ).all()

    # 拉取该学生的多维度数据
    student_context = None
    student = c.student if c else None
    if student:
        student_context = _build_student_context(student)

    return render_template("endterm_comment/form.html",
                           students=students,
                           comment=c,
                           student=student,
                           student_context=student_context,
                           pre_student_id=c.student_id,
                           semester=c.semester)


@endterm_comment_bp.route("/<int:cid>/delete", methods=["POST"])
@login_required
def delete(cid):
    c = EndTermComment.query.get_or_404(cid)
    if session.get("role") == "class_teacher" and c.class_id != session.get("class_id"):
        flash("无权操作", "danger")
        return redirect(url_for("endterm_comment.index"))
    db.session.delete(c)
    safe_commit()
    flash("已删除", "info")
    return redirect(url_for("endterm_comment.index"))


@endterm_comment_bp.route("/batch", methods=["GET", "POST"])
@login_required
@require_role("class_teacher")
def batch():
    if request.method == "POST":
        return jsonify({"ok": True})
    return render_template("endterm_comment/batch.html")


# ══════════════════════════════════════════════════════════════
#  🚀 AI 评语生成舱 — 单生 + 批量端点
# ══════════════════════════════════════════════════════════════

@endterm_comment_bp.route("/api/ai-generate/<int:student_id>", methods=["POST"])
@login_required
@require_role("class_teacher", "ms_admin", "grade_leader")
def ai_generate_comment(student_id):
    """为大模型生成单个学生评语 — POST请求触发LLM调用"""
    stu = Student.query.get_or_404(student_id)

    # 班主任数据隔离
    user_role = session.get("role")
    if user_role == "class_teacher":
        teacher_class = session.get("class_id")
        if teacher_class and stu.class_id != teacher_class:
            return jsonify({"status": "error", "message": "无权访问该学生"}), 403

    # 1. 构建增强版数字灵魂画像
    context = _build_llm_context(stu)

    # 2. 序列化为LLM友好格式
    user_prompt = _format_llm_prompt(stu, context)

    # 3. 调用大模型
    try:
        ai_raw = _call_llm_api(COMMENT_EXPERT_SYSTEM_PROMPT, user_prompt)
        comment_data = json.loads(ai_raw)

        # 验证必需字段
        required = ["strengths", "improvements", "overall_comment", "teacher_suggestion"]
        for key in required:
            if key not in comment_data:
                comment_data[key] = ""

        return jsonify({
            "status": "success",
            "data": comment_data,
            "student_name": stu.name,
            "student_id": stu.id
        })

    except json.JSONDecodeError:
        current_app.logger.error(f"LLM返回非JSON格式: {ai_raw[:200] if 'ai_raw' in dir() else 'N/A'}")
        return jsonify({"status": "error", "message": "大模型返回格式异常，请重试"}), 500
    except requests.exceptions.Timeout:
        return jsonify({"status": "error", "message": "大模型响应超时，请稍后重试"}), 504
    except requests.exceptions.RequestException as e:
        current_app.logger.error(f"LLM API调用失败: {str(e)}")
        return jsonify({"status": "error", "message": f"大模型服务暂不可用: {str(e)[:100]}"}), 502
    except Exception as e:
        current_app.logger.error(f"评语生成未知错误: {str(e)}")
        return jsonify({"status": "error", "message": "评语织网机遭遇未知故障"}), 500


@endterm_comment_bp.route("/api/ai-batch-generate", methods=["POST"])
@login_required
@require_role("class_teacher")
def ai_batch_generate():
    """全班批量生成评语 — 返回学生列表+生成状态（异步逐条处理）"""
    class_id = session.get("class_id")
    if not class_id:
        return jsonify({"status": "error", "message": "无法获取您的班级"}), 400

    students = Student.query.filter_by(class_id=class_id, is_active=True).order_by(
        Student.student_no
    ).all()

    if not students:
        return jsonify({"status": "error", "message": "班级无在读学生"}), 404

    results = []
    success_count = 0
    fail_count = 0

    for stu in students:
        try:
            context = _build_llm_context(stu)
            user_prompt = _format_llm_prompt(stu, context)
            ai_raw = _call_llm_api(COMMENT_EXPERT_SYSTEM_PROMPT, user_prompt)
            comment_data = json.loads(ai_raw)
            results.append({
                "student_id": stu.id,
                "student_name": stu.name,
                "status": "success",
                "data": comment_data
            })
            success_count += 1
        except Exception as e:
            results.append({
                "student_id": stu.id,
                "student_name": stu.name,
                "status": "error",
                "message": str(e)[:100]
            })
            fail_count += 1

    return jsonify({
        "status": "success" if fail_count == 0 else "partial",
        "total": len(students),
        "success": success_count,
        "failed": fail_count,
        "results": results
    })


def _format_llm_prompt(stu, context):
    """将结构化数据编织为LLM的自然语言提示"""
    parts = []
    parts.append(f"学生姓名：{stu.name}")
    parts.append(f"班级：{stu.class_.name if stu.class_ else '未知'}")
    parts.append(f"性别：{stu.gender or '未知'}")

    # 成绩斜率
    if context.get("score_trends"):
        parts.append("\n【考试成绩趋势】")
        for subj, info in context["score_trends"].items():
            direction = "↑" if info["slope"] > 0 else ("↓" if info["slope"] < 0 else "→")
            parts.append(
                f"  {subj}: {info['first_score']}→{info['last_score']} "
                f"({direction}{info['slope']}分/次, {info['exams_count']}次考试, 整体{info['trend']})"
            )

    # 最近考试成绩
    if context.get("scores") and context["scores"].get("exam_name"):
        sc = context["scores"]
        parts.append(f"\n【最近考试】{sc['exam_name']} ({sc['exam_date']})")
        parts.append(f"  均分: {sc['avg']}")
        for s in sc.get("scores", []):
            subj_name = s.subject.name if s.subject else "?"
            parts.append(f"  {subj_name}: {s.score}分")

    # 五翼行为分
    if context.get("wings"):
        parts.append("\n【五维行为分沉淀】")
        dim_names = {"德": "品德修养", "智": "学业表现", "体": "身心健康", "美": "审美素养", "劳": "劳动实践"}
        for dim, info in context["wings"].items():
            dim_label = dim_names.get(dim, dim)
            parts.append(f"  {dim_label}({dim}): 均分{info['avg']}分, 共{info['count']}次评分")

    # 违纪
    disc = context.get("discipline", {})
    if disc.get("total", 0) > 0:
        parts.append(f"\n【违纪记录】共{disc['total']}条, 累计扣分{disc.get('points', 0)}分")

    # 活动
    acts = context.get("activities", {})
    if acts.get("total", 0) > 0:
        act_names = [a.title for a in acts.get("list", []) if hasattr(a, 'title')]
        parts.append(f"\n【活动参与】共{acts['total']}次: {', '.join(act_names[:3])}")

    # 心理健康
    mh = context.get("mental_health", {})
    if mh.get("risk_level"):
        parts.append(f"\n【心理健康】风险等级: {mh['risk_level']}, 总分: {mh.get('total_score', '?')}")

    # 班主任手记
    if context.get("teacher_notes"):
        parts.append("\n【班主任随访手记】")
        for note in context["teacher_notes"]:
            parts.append(f"  [{note['category']}] {note['content'][:150]}")

    # 干预效果
    if context.get("interventions"):
        parts.append("\n【历史干预效果】")
        for iv in context["interventions"]:
            parts.append(f"  {iv['type']}: {iv['effect']}")

    return "\n".join(parts)


@endterm_comment_bp.route("/export")
@login_required
def export():
    """导出期末评语 — Word / HTML打印"""
    semester = request.args.get("semester", "")
    class_id = request.args.get("class_id", type=int)
    type_    = request.args.get("type", "word")      # word | pdf(html)

    # 构造查询
    q = EndTermComment.query
    # 班主任只能导出本班评语
    if session.get("role") == "class_teacher":
        class_id = session.get("class_id")
    if class_id:
        q = q.join(Student).filter(Student.class_id == class_id)
    comments = q.order_by(EndTermComment.semester.desc(),
                        EndTermComment.id.asc()).all()

    if not comments:
        flash("暂无评语数据可导出", "warning")
        return redirect(url_for("endterm_comment.index"))

    if type_ == "word":
        return _export_word(comments, semester)
    else:
        return _export_html(comments, semester)


# ── Word 导出 ──────────────────────────────
def _export_word(comments, semester):
    """生成 .docx 并返回"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        flash("服务器未安装 python-docx，请联系管理员安装", "danger")
        return redirect(url_for("endterm_comment.index"))

    doc = Document()

    # 标题
    h = doc.add_heading("梨江中学 期末评语", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"学期：{semester or '全部'}")
    doc.add_paragraph("")

    for c in comments:
        s = c.student
        p = doc.add_paragraph()
        p.add_run(f"【{s.name if s else '?'}】  ").bold = True
        if s and s.class_:
            p.add_run(f"{s.class_.name}  ")
        doc.add_paragraph(f"综合评语：{c.overall_comment or '—'}")
        if c.strengths:
            doc.add_paragraph(f"主要优点：{c.strengths}")
        if c.improvements:
            doc.add_paragraph(f"改进建议：{c.improvements}")
        doc.add_paragraph("—" * 30)

    # 保存到内存
    import io
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    from flask import make_response
    resp = make_response(buf.read())
    resp.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    resp.headers["Content-Disposition"] = f"attachment; filename=期末评语_{semester or '全部'}.docx"
    return resp


# ── HTML 打印版（可 Ctrl+P 保存PDF）─
def _export_html(comments, semester):
    from flask import make_response
    html = render_template("endterm_comment/print.html",
                          comments=comments, semester=semester)
    resp = make_response(html)
    resp.headers["Content-Type"] = "text/html; charset=utf-8"
    return resp


# ── 学生多维度数据聚合 ──────────────────────────────
# 缓存最近一次考试（全校通用，1分钟 TTL）—— 避免逐学生重复查询
_latest_exam_cache = {"value": None, "ts": 0}
_LATEST_EXAM_CACHE_TTL = 60  # 秒


def _get_latest_exam():
    now = time.time()
    if now - _latest_exam_cache["ts"] < _LATEST_EXAM_CACHE_TTL:
        return _latest_exam_cache["value"]
    from models import Exam
    exam = Exam.query.order_by(Exam.exam_date.desc()).first()
    _latest_exam_cache["value"] = exam
    _latest_exam_cache["ts"] = now
    return exam


def _build_student_context(stu):
    """拉取学生多维度数据，供期末评语参考"""
    from models import DisciplineRecord, Score, Exam, ActivityRegistration, Activity, MentalHealthAssessment
    from sqlalchemy import func
    from sqlalchemy.orm import joinedload

    ctx = {"student": stu, "discipline": {}, "scores": {}, "activities": {}, "mental_health": {}}

    # 1. 违纪记录统计
    records = DisciplineRecord.query.filter_by(student_id=stu.id).all()
    ctx["discipline"] = {
        "total": len(records),
        "active": sum(1 for r in records if r.status == "active"),
        "points": sum(r.points for r in records),
        "recent": records[-3:] if records else [],
    }

    # 2. 最近考试成绩 — latest_exam 缓存避免逐学生重复查
    latest_exam = _get_latest_exam()
    if latest_exam:
        scores = Score.query.filter_by(student_id=stu.id, exam_id=latest_exam.id).all()
        ctx["scores"] = {
            "exam_name": latest_exam.name,
            "exam_date": latest_exam.exam_date,
            "scores": scores,
            "avg": round(sum(s.score for s in scores) / len(scores), 1) if scores else 0,
        }

    # 3. 活动参与 — joinedload 消除循环内 N+1 Activity 查询
    regs = ActivityRegistration.query.filter_by(
        student_id=stu.id, status="confirmed"
    ).options(joinedload(ActivityRegistration.activity)).all()
    activity_list = [reg.activity for reg in regs if reg.activity]
    ctx["activities"] = {
        "total": len(regs),
        "list": activity_list[:5],
    }

    # 4. 心理健康评估
    mh = MentalHealthAssessment.query.filter_by(student_id=stu.id).order_by(
        MentalHealthAssessment.created_at.desc()
    ).first()
    if mh:
        ctx["mental_health"] = {
            "total_score": mh.total_score,
            "risk_level": mh.risk_level,
            "conclusion": mh.conclusion,
        }

    return ctx


def _build_llm_context(stu):
    """为 LLM 构建增强版学生数字灵魂画像（含成绩斜率、行为分趋势、班主任手记）"""
    ctx = _build_student_context(stu)
    from models import Subject

    # ── 5. 考试斜率分析（所有考试，按日期排序）──
    exams = Exam.query.filter_by(grade_id=stu.grade_id).order_by(Exam.exam_date.asc()).all()
    subjects = Subject.query.order_by(Subject.sort_order).all()
    subject_map = {s.id: s.name for s in subjects}

    if len(exams) >= 2 and subjects:
        all_scores = Score.query.filter(
            Score.student_id == stu.id,
            Score.exam_id.in_([e.id for e in exams])
        ).all()

        # 构建成绩时间线: {subject_name: [(exam_name, score), ...]}
        exam_map = {e.id: e.name for e in exams}
        timeline = {}
        for s in all_scores:
            subj_name = subject_map.get(s.subject_id, f"科目{s.subject_id}")
            if subj_name not in timeline:
                timeline[subj_name] = []
            timeline[subj_name].append((exam_map.get(s.exam_id, "?"), s.score))

        # 计算各科斜率
        slopes = {}
        for subj, records in timeline.items():
            if len(records) >= 2:
                records.sort(key=lambda x: x[0])
                first, last = records[0], records[-1]
                n = len(records)
                slope = round((last[1] - first[1]) / max(n - 1, 1), 1)
                slopes[subj] = {
                    "slope": slope,
                    "trend": "上升" if slope > 0 else ("下降" if slope < 0 else "平稳"),
                    "first_score": first[1],
                    "last_score": last[1],
                    "exams_count": n
                }
        ctx["score_trends"] = slopes

    # ── 6. 五翼行为分汇总 ──
    wing_scores = WingsScore.query.filter_by(student_id=stu.id).all()
    if wing_scores:
        dim_summary = {}
        for ws in wing_scores:
            d = ws.dimension
            if d not in dim_summary:
                dim_summary[d] = {"total": 0, "count": 0, "recent": []}
            dim_summary[d]["total"] += ws.score
            dim_summary[d]["count"] += 1
            dim_summary[d]["recent"].append({
                "score": ws.score,
                "scorer_type": ws.scorer_type,
                "date": ws.created_at.strftime("%m-%d") if ws.created_at else "?"
            })
        for d in dim_summary:
            dim_summary[d]["avg"] = round(dim_summary[d]["total"] / dim_summary[d]["count"], 1)
            dim_summary[d]["recent"] = dim_summary[d]["recent"][-5:]  # 最近5条
        ctx["wings"] = dim_summary

    # ── 7. 班主任手记（最近3条）──
    notes = TeacherNote.query.filter_by(student_id=stu.id).order_by(
        TeacherNote.created_at.desc()
    ).limit(3).all()
    if notes:
        ctx["teacher_notes"] = [
            {"category": n.category, "content": n.content, "date": n.created_at.strftime("%m-%d") if n.created_at else "?"}
            for n in notes
        ]

    # ── 8. 干预效果 ──
    interventions = InterventionRecord.query.filter_by(
        student_id=stu.id
    ).order_by(InterventionRecord.created_at.desc()).limit(5).all()
    if interventions:
        ctx["interventions"] = [
            {
                "type": iv.intervention_type,
                "effect": iv.effect_rating,
                "notes": iv.notes[:100] if iv.notes else ""
            }
            for iv in interventions
        ]

    return ctx

